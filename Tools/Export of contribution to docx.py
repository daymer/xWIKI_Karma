from docx import Document
from docx.enum.text import WD_COLOR_INDEX

import Configuration
from CustomModules.Mechanics import XWikiClient
from CustomModules.SQL_Connector import SQLConnector

SQLConfig = Configuration.SQLConfig()
SQLConnector = SQLConnector(SQLConfig)
xWikiConfig = Configuration.XWikiConfig('Migration pool')
xWikiClient = XWikiClient(xWikiConfig.api_root, xWikiConfig.auth_user, xWikiConfig.auth_pass)

PageTitle = 'Using WMI to query Veeam BnR information'
platform = 'xWIKI'
SQLQuery = SQLConnector.select_datagrams_from_dbo_knownpages_datagrams(PageTitle, platform)
datagram = SQLQuery[0]
contributors_datagram = SQLQuery[1]
Colors = ['BRIGHT_GREEN','YELLOW','TEAL','VIOLET','PINK','RED','TURQUOISE','DARK_YELLOW','GRAY_50','GRAY_25','DARK_BLUE','DARK_RED','BLUE','GREEN']
ColorsByUser = {}
UniqueUsers = set(contributors_datagram.values())
for index, contributor in enumerate(UniqueUsers):
    ColorsByUser.update({contributor: Colors[index]})
#print(dir(WD_COLOR_INDEX))
#making docx
document = Document()
document.add_heading(PageTitle)
#filling docx
paragraph_authors = document.add_paragraph()
paragraph_authors.add_run('Legend:')
run = paragraph_authors.add_run()
run.add_break()
for user in ColorsByUser:
    Current_color = ColorsByUser[user]
    font = paragraph_authors.add_run(user).font
    font.highlight_color = getattr(WD_COLOR_INDEX, Current_color)
    run = paragraph_authors.add_run()
    run.add_break()

paragraph = document.add_paragraph()
paragraph.add_run('Page content by authors:')
run = paragraph.add_run()
run.add_break()
for letter in datagram:
    Current_color = ColorsByUser[contributors_datagram[letter[1]]]
    font = paragraph.add_run(letter[0]).font
    font.highlight_color = getattr(WD_COLOR_INDEX,Current_color)
document.save('C:\Temp\\' +PageTitle+ '.docx')
